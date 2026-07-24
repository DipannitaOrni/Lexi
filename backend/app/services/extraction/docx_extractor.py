"""
DOCX extraction via python-docx.

Body paragraphs and tables are both read — tables are flattened into a
readable "row: col1 | col2 | col3" text form so tabular data isn't silently
dropped from the pipeline (a document's key facts are often in a table).
"""
import io

from docx import Document

from app.services.extraction.base import ExtractedDocument


def _flatten_table(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_docx(file_bytes: bytes) -> ExtractedDocument:
    doc = Document(io.BytesIO(file_bytes))

    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        flattened = _flatten_table(table)
        if flattened.strip():
            parts.append(flattened)

    raw_text = "\n\n".join(parts)
    return ExtractedDocument(
        raw_text=raw_text,
        page_count=1,  # DOCX has no reliable page concept without rendering
        warnings=[],
        source_format="docx",
    )
